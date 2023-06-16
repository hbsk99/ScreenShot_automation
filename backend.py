import pandas as pd
from pynput import keyboard, mouse
from itertools import zip_longest
import pyautogui
from docx import Document
from docx.shared import Inches
from io import BytesIO

keyboard_events = []
mouse_events = []
screenshots = []
keyboard_listener = None
mouse_listener = None

def on_press(key):
    keyboard_events.append(f"Key pressed: {key}")

def on_click(x, y, button, pressed):
    if pressed:
        screenshot = pyautogui.screenshot()
        screenshots.append(screenshot)
        mouse_events.append(f"Mouse clicked at ({x}, {y}) with button {button}")

def capture_events(status_label, stop=False):
    global keyboard_listener, mouse_listener

    if stop:
        if keyboard_listener and mouse_listener:
            keyboard_listener.stop()
            mouse_listener.stop()
            update_status_label(status_label, 'Screen events capturing stopped!')
            save_data_to_excel()
            save_data_to_word()
    else:
        keyboard_events.clear()
        mouse_events.clear()
        screenshots.clear()
        start_capture()
        update_status_label(status_label, 'Screen events capturing started!')

def start_capture():
    global keyboard_listener, mouse_listener

    keyboard_listener = keyboard.Listener(on_press=on_press)
    mouse_listener = mouse.Listener(on_click=on_click)

    keyboard_listener.start()
    mouse_listener.start()

def start_program(status_label):
    capture_events(status_label)

def update_status_label(status_label, text):
    # Update the status label in the UI with the provided text
    status_label.config(text=text)

def save_data_to_excel():
    data = zip_longest(keyboard_events, mouse_events, screenshots, fillvalue='')

    df = pd.DataFrame(data, columns=['Keyboard Events', 'Mouse Events', 'Screenshots'])

    with pd.ExcelWriter('screen_events.xlsx') as writer:
        df.to_excel(writer, index=False)

    save_screenshots_to_word()
   
def save_screenshots_to_word():
    document = Document()

    for screenshot in screenshots:
        image_stream = BytesIO()
        screenshot.save(image_stream, format='PNG')
        image_stream.seek(0)

        document.add_picture(image_stream, width=Inches(6))

    document.save('screenshots.docx')


def save_data_to_word():
    # Create a Word document
    doc = Document()

    # Add the captured events to the Word document
    doc.add_heading('Captured Events', level=1)

    # Add keyboard events
    doc.add_heading('Keyboard Events', level=2)
    for event in keyboard_events:
        doc.add_paragraph(event)

    # Add mouse events and screenshots
    doc.add_heading('Mouse Events', level=2)
    for event, screenshot in zip(mouse_events, screenshots):
        doc.add_paragraph(event)

        # Convert screenshot to file-like object
        image_stream = BytesIO()
        screenshot.save(image_stream, format='PNG')
        image_stream.seek(0)

        # Add the screenshot to the document
        doc.add_picture(image_stream, width=Inches(6))

    # Save the Word document
    doc.save('captured_events.docx')
